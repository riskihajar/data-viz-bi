from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
ARTICLE_DIR = ROOT / "docs" / "artikel-ieee"
OUTPUT_DOCX = ARTICLE_DIR / "Artikel IEEE - Early Warning OULAD.docx"
OUTPUT_MD = ARTICLE_DIR / "artikel-ieee.md"

SECTION_FILES = [
    "00-title-author.md",
    "01-abstract-keywords.md",
    "02-introduction.md",
    "03-related-works.md",
    "04-methodology.md",
    "05-results.md",
    "06-discussion.md",
    "07-conclusion.md",
    "09-references.md",
]

AUTHORS = [
    ("Muhammad Rizky Hajar", "riskihajar@students.amikom.ac.id"),
    ("Alwie Muflich", "alwiemuflich@students.amikom.ac.id"),
    ("Heri Santosa", "heri.sant@students.amikom.ac.id"),
    ("Andi Sunyoto", "andi@amikom.ac.id"),
    ("Robert Marco", "robert.marco@amikom.ac.id"),
]

FULL_WIDTH_FIGURES = {
    "fig-4-dashboard-dvbi.png",
}

FIGURE_WIDTHS = {
    "fig-1a-target-distribution.png": 3.25,
    "fig-1b-missing-values.png": 3.25,
    "fig-2a-metrics-comparison.png": 3.25,
    "fig-2b-confusion-matrix.png": 3.25,
    "fig-2c-roc-curve.png": 3.25,
    "fig-3-feature-importance.png": 3.25,
    "fig-4-dashboard-dvbi.png": 6.0,
}


def set_cell_margins(cell, top=60, start=80, bottom=60, end=80):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths_dxa):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_columns(section, count: int, space_twips: int = 360):
    sect_pr = section._sectPr
    cols = sect_pr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sect_pr.append(cols)
    cols.set(qn("w:num"), str(count))
    cols.set(qn("w:space"), str(space_twips))
    cols.set(qn("w:equalWidth"), "1")


def set_run_font(run, size=10, bold=None, italic=None):
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 0, 0)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_inline_markdown(paragraph, text: str, size=10, bold=False, italic=False):
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*|\*[^*]+\*)", text)
    for part in parts:
        if not part:
            continue
        run_bold, run_italic = bold, italic
        content = part
        if part.startswith("**") and part.endswith("**"):
            content = part[2:-2]
            run_bold = True
        elif part.startswith("*") and part.endswith("*"):
            content = part[1:-1]
            run_italic = True
        elif part.startswith("`") and part.endswith("`"):
            content = part[1:-1]
            run_italic = True
        run = paragraph.add_run(content)
        set_run_font(run, size=size, bold=run_bold, italic=run_italic)


def add_body_paragraph(doc, text: str, italic=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.18)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.0
    add_inline_markdown(p, text, size=10, italic=italic)
    return p


def add_main_heading(doc, text: str):
    p = doc.add_paragraph(style="Heading 1")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text.upper())
    set_run_font(run, size=10, bold=False)
    p.paragraph_format.keep_with_next = True


def add_subheading(doc, text: str):
    p = doc.add_paragraph(style="Heading 2")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_run_font(run, size=10, italic=True)
    p.paragraph_format.keep_with_next = True


def parse_table(lines, start):
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", c) for c in cells):
            rows.append(cells)
        i += 1
    return rows, i


def add_table_caption(doc, caption: str):
    match = re.match(r"(?:Table|Tabel)\s+([IVXLCDM]+)\.\s*(.+)", caption, re.IGNORECASE)
    number, title = (match.group(1), match.group(2)) if match else ("", caption)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    if number:
        run = p.add_run(f"TABLE {number.upper()}\n")
        set_run_font(run, size=8)
        run = p.add_run(title.upper())
        set_run_font(run, size=8)
    else:
        add_inline_markdown(p, title, size=8)


def add_result_table(doc, rows, caption=None):
    if not rows:
        return
    compact = len(rows[0]) <= 4
    current = doc.sections[-1]
    if not compact:
        one_col = doc.add_section(WD_SECTION.CONTINUOUS)
        one_col.top_margin = current.top_margin
        one_col.bottom_margin = current.bottom_margin
        one_col.left_margin = current.left_margin
        one_col.right_margin = current.right_margin
        set_columns(one_col, 1)

    if caption:
        add_table_caption(doc, caption)

    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    total = 4800 if compact else 10260
    if compact:
        first = 1100 if len(rows[0]) == 4 else 1400
    else:
        first = 2600
    remaining = total - first
    widths = [first] + [remaining // (len(rows[0]) - 1)] * (len(rows[0]) - 1)
    widths[-1] += total - sum(widths)
    set_table_widths(table, widths)
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            add_inline_markdown(p, value, size=6.8 if compact else 7.2, bold=(r_idx == 0))
            if r_idx == 0:
                shading = OxmlElement("w:shd")
                shading.set(qn("w:fill"), "E7E6E6")
                cell._tc.get_or_add_tcPr().append(shading)
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))

    if not compact:
        two_col = doc.add_section(WD_SECTION.CONTINUOUS)
        two_col.top_margin = current.top_margin
        two_col.bottom_margin = current.bottom_margin
        two_col.left_margin = current.left_margin
        two_col.right_margin = current.right_margin
        set_columns(two_col, 2)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(7)
    spacer.paragraph_format.line_spacing = 1.0


def add_result_figure(doc, image_path: Path, caption: str):
    current = doc.sections[-1]
    is_full_width = image_path.name in FULL_WIDTH_FIGURES
    if is_full_width:
        one_col = doc.add_section(WD_SECTION.CONTINUOUS)
        one_col.top_margin = current.top_margin
        one_col.bottom_margin = current.bottom_margin
        one_col.left_margin = current.left_margin
        one_col.right_margin = current.right_margin
        set_columns(one_col, 1)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.keep_with_next = True
    shape = p.add_run().add_picture(
        str(image_path), width=Inches(FIGURE_WIDTHS.get(image_path.name, 3.25))
    )
    doc_pr = shape._inline.docPr
    doc_pr.set("descr", caption)
    doc_pr.set("title", caption.split(".", 1)[0])

    caption_p = doc.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_p.paragraph_format.space_before = Pt(1)
    caption_p.paragraph_format.space_after = Pt(4)
    caption_p.paragraph_format.keep_together = True
    add_inline_markdown(caption_p, caption, size=8)

    if is_full_width:
        two_col = doc.add_section(WD_SECTION.CONTINUOUS)
        two_col.top_margin = current.top_margin
        two_col.bottom_margin = current.bottom_margin
        two_col.left_margin = current.left_margin
        two_col.right_margin = current.right_margin
        set_columns(two_col, 2)


def add_markdown_section(doc, path: Path):
    lines = path.read_text(encoding="utf-8").splitlines()
    i = 0
    paragraph_buffer = []
    pending_table_caption = None

    def flush():
        nonlocal paragraph_buffer
        if paragraph_buffer:
            add_body_paragraph(doc, " ".join(x.strip() for x in paragraph_buffer))
            paragraph_buffer = []

    while i < len(lines):
        line = lines[i].rstrip()
        if not line:
            flush()
            i += 1
            continue
        if line.startswith("# "):
            flush()
            add_main_heading(doc, line[2:].strip())
            i += 1
            continue
        if line.startswith("## "):
            flush()
            add_subheading(doc, line[3:].strip())
            i += 1
            continue
        if line.startswith("**Table ") or line.startswith("**Tabel "):
            flush()
            pending_table_caption = line.strip("*")
            i += 1
            continue
        if line.startswith("|"):
            flush()
            rows, i = parse_table(lines, i)
            add_result_table(doc, rows, pending_table_caption)
            pending_table_caption = None
            continue
        figure = re.fullmatch(r"!\[(.+)]\((.+)\)", line)
        if figure:
            flush()
            caption, relative_path = figure.groups()
            add_result_figure(doc, path.parent / relative_path, caption)
            i += 1
            continue
        if re.match(r"^\[\d+\]", line):
            flush()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.left_indent = Inches(0.18)
            p.paragraph_format.first_line_indent = Inches(-0.18)
            p.paragraph_format.space_after = Pt(2)
            add_inline_markdown(p, line, size=8)
            i += 1
            continue
        paragraph_buffer.append(line)
        i += 1
    flush()


def build_combined_markdown():
    chunks = [f"# Artikel IEEE Lengkap\n"]
    for name in SECTION_FILES:
        chunks.append((ARTICLE_DIR / name).read_text(encoding="utf-8").strip())
    OUTPUT_MD.write_text("\n\n".join(chunks) + "\n", encoding="utf-8")


def build_docx():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)
    set_columns(section, 1)

    title = (ARTICLE_DIR / "00-title-author.md").read_text(encoding="utf-8").splitlines()[0][2:]
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(title)
    set_run_font(run, size=20)

    author_table = doc.add_table(rows=2, cols=3)
    author_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    author_table.autofit = False
    set_table_widths(author_table, [3420, 3420, 3420])
    author_table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    author_cells = [cell for row in author_table.rows for cell in row.cells]
    for cell in author_cells:
        cell._tc.get_or_add_tcPr().append(OxmlElement("w:tcBorders"))
    for cell, (name, email) in zip(author_cells, AUTHORS):
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        lines = [
            (name, 10, False),
            ("\nDepartment of Computer Science", 9, False),
            ("\nUniversitas Amikom Yogyakarta", 9, False),
            ("\nYogyakarta, Indonesia", 9, False),
        ]
        if email:
            lines.append((f"\n{email}", 8.5, True))
        for text, size, italic in lines:
            r = p.add_run(text)
            set_run_font(r, size=size, italic=italic)
    for cell in author_cells[len(AUTHORS) :]:
        cell.text = ""

    body_section = doc.add_section(WD_SECTION.CONTINUOUS)
    body_section.top_margin = Inches(0.65)
    body_section.bottom_margin = Inches(0.65)
    body_section.left_margin = Inches(0.65)
    body_section.right_margin = Inches(0.65)
    set_columns(body_section, 2)

    abstract_lines = (ARTICLE_DIR / "01-abstract-keywords.md").read_text(encoding="utf-8").splitlines()
    abstract = next(x for x in abstract_lines[1:] if x.strip())
    keyword_idx = abstract_lines.index("# Keywords")
    keywords = next(x for x in abstract_lines[keyword_idx + 1 :] if x.strip())
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Abstract. ")
    set_run_font(r, size=9, bold=True, italic=True)
    add_inline_markdown(p, abstract, size=9)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run("Keywords. ")
    set_run_font(r, size=9, bold=True, italic=True)
    add_inline_markdown(p, keywords, size=9)

    for name in SECTION_FILES[2:]:
        add_markdown_section(doc, ARTICLE_DIR / name)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(10)

    doc.core_properties.title = title
    doc.core_properties.author = ", ".join(name for name, _ in AUTHORS)
    doc.core_properties.subject = "Early warning student dropout using OULAD"
    doc.core_properties.keywords = "student dropout, OULAD, early warning, knowledge-based system, business intelligence"
    doc.save(OUTPUT_DOCX)


if __name__ == "__main__":
    build_combined_markdown()
    build_docx()
    print(f"Wrote {OUTPUT_MD}")
    print(f"Wrote {OUTPUT_DOCX}")
